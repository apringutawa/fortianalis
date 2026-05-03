"""
Upload & Export API Router
Endpoints:
  POST /api/v1/upload/           — upload & analyze a FortiWeb log file
  GET  /api/v1/upload/export/{file_id}  — export PDF report
  GET  /api/v1/upload/export/{file_id}/excel  — export Excel report
  GET  /api/v1/upload/export/{file_id}/word   — export Word report
  GET  /api/v1/upload/export/{file_id}/csv    — export CSV report
"""
import os
import uuid
import csv as csv_module
import io
from datetime import datetime, timedelta

from fastapi import APIRouter, UploadFile, File, HTTPException, Query, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.db import get_db
from app.models import Report
from app.utils.parser import parse_fortiweb_log
from app.utils.ai import generate_security_insight, generate_report_narrative

router = APIRouter()

UPLOAD_DIR = "data/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
#  Upload & Analysis
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/")
async def upload_log_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    allowed_exts = (".log", ".txt", ".csv")
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in allowed_exts:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type '{ext}'. Allowed: {', '.join(allowed_exts)}",
        )

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File exceeds 50 MB limit.")

    file_id = str(uuid.uuid4())
    saved_path = os.path.join(UPLOAD_DIR, f"{file_id}{ext}")

    with open(saved_path, "wb") as fh:
        fh.write(content)

    try:
        parsed_data = parse_fortiweb_log(saved_path)
        ai_insight = generate_security_insight(parsed_data)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Processing error: {exc}")

    stats = parsed_data.get("stats", {})
    is_demo = parsed_data.get("_demo", False)

    db_report = Report(
        id=file_id,
        filename=file.filename or "unknown",
        created_at=datetime.utcnow(),
        total_requests=stats.get("totalRequests", 0),
        total_attacks=stats.get("totalAttacks", 0),
        blocked_attacks=stats.get("blockedAttacks", 0),
        unique_ips=stats.get("uniqueIps", 0),
        block_rate=stats.get("blockRate", 0.0),
        attack_types=parsed_data.get("attackTypes", []),
        attacker_ips=parsed_data.get("attackerIPs", []),
        subdomains=parsed_data.get("subdomains", []),
        timeline_data=parsed_data.get("timelineData", []),
        ai_analysis=ai_insight.get("analysis", ""),
        ai_recommendation=ai_insight.get("recommendation", ""),
        ai_powered_by=ai_insight.get("powered_by", ""),
        is_demo=1 if is_demo else 0,
    )
    db.add(db_report)
    db.commit()

    result = {
        "message": "File processed successfully",
        "file_id": file_id,
        "filename": file.filename,
        "data": parsed_data,
        "ai_insight": ai_insight,
        "is_demo": is_demo,
    }
    return result


# ─────────────────────────────────────────────────────────────────────────────
#  Multi-Format Export
# ─────────────────────────────────────────────────────────────────────────────

@router.get("/export/{file_id}")
async def export_pdf(
    file_id: str,
    period: str = Query(default="daily", enum=["daily", "weekly", "monthly"]),
    db: Session = Depends(get_db),
):
    report = db.query(Report).filter(Report.id == file_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")

    # Generate comprehensive narrative
    parsed_data = {
        "stats": {
            "totalRequests": report.total_requests,
            "totalAttacks": report.total_attacks,
            "blockedAttacks": report.blocked_attacks,
            "uniqueIps": report.unique_ips,
            "blockRate": report.block_rate,
        },
        "attackTypes": report.attack_types or [],
        "attackerIPs": report.attacker_ips or [],
        "subdomains": report.subdomains or [],
        "timelineData": report.timeline_data or [],
    }
    ai_insight = {
        "analysis": report.ai_analysis,
        "recommendation": report.ai_recommendation,
        "powered_by": report.ai_powered_by,
    }
    narrative = generate_report_narrative(parsed_data, ai_insight)

    pdf_bytes = _generate_pdf(report, period, narrative)
    period_label = {"daily": "Daily", "weekly": "Weekly", "monthly": "Monthly"}[period]
    filename = f"FortiAnalis_{period_label}_Report_{datetime.now().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/export/{file_id}/excel")
async def export_excel(
    file_id: str,
    db: Session = Depends(get_db),
):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.chart.line_chart import LineChart
        from openpyxl.chart.reference import Reference
        from openpyxl.chart.series import DataPoint
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise HTTPException(status_code=501, detail="openpyxl is not installed.")

    report = db.query(Report).filter(Report.id == file_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")

    # Generate comprehensive narrative
    parsed_data = {
        "stats": {
            "totalRequests": report.total_requests,
            "totalAttacks": report.total_attacks,
            "blockedAttacks": report.blocked_attacks,
            "uniqueIps": report.unique_ips,
            "blockRate": report.block_rate,
        },
        "attackTypes": report.attack_types or [],
        "attackerIPs": report.attacker_ips or [],
        "subdomains": report.subdomains or [],
        "timelineData": report.timeline_data or [],
    }
    ai_insight = {
        "analysis": report.ai_analysis,
        "recommendation": report.ai_recommendation,
        "powered_by": report.ai_powered_by,
    }
    narrative = generate_report_narrative(parsed_data, ai_insight)

    wb = Workbook()

    title_font = Font(name="Calibri", size=16, bold=True, color="06B6D4")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
    row_fill_1 = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    row_fill_2 = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin", color="E2E8F0"),
        right=Side(style="thin", color="E2E8F0"),
        top=Side(style="thin", color="E2E8F0"),
        bottom=Side(style="thin", color="E2E8F0"),
    )

    # Sheet 1: Executive Summary
    ws1 = wb.active
    ws1.title = "Executive Summary"
    ws1.sheet_properties.tabColor = "06B6D4"

    ws1.merge_cells("A1:B1")
    ws1["A1"] = "FortiAnalis - WAF Security Report"
    ws1["A1"].font = title_font

    ws1["A2"] = f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}"
    ws1["A2"].font = Font(name="Calibri", size=10, color="64748B")
    ws1["A3"] = f"File: {report.filename}"
    ws1["A3"].font = Font(name="Calibri", size=10, color="64748B")

    # Executive Summary Narrative
    ws1["A5"] = "Executive Summary"
    ws1["A5"].font = Font(name="Calibri", size=13, bold=True, color="94A3B8")
    exec_summary = narrative.get("executive_summary", "")
    ws1.merge_cells("A6:B10")
    ws1["A6"] = exec_summary
    ws1["A6"].font = Font(name="Calibri", size=10)
    ws1["A6"].alignment = Alignment(wrap_text=True, vertical="top")
    ws1.row_dimensions[6].height = 120

    summary_data = [
        ["Metric", "Value"],
        ["Total Requests", report.total_requests],
        ["Total Attacks", report.total_attacks],
        ["Blocked Attacks", report.blocked_attacks],
        ["Block Rate", f"{report.block_rate}%"],
        ["Unique Attacker IPs", report.unique_ips],
    ]
    for i, row in enumerate(summary_data, start=12):
        for j, val in enumerate(row, start=1):
            cell = ws1.cell(row=i, column=j, value=val)
            cell.border = thin_border
            if i == 12:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
            else:
                cell.fill = row_fill_1 if i % 2 == 0 else row_fill_2
        ws1.column_dimensions["A"].width = 25
        ws1.column_dimensions["B"].width = 20
    ws1.row_dimensions[14].height = 60

    # Sheet 2: Attack Types
    ws2 = wb.create_sheet("Attack Types")
    ws2.sheet_properties.tabColor = "3B82F6"
    ws2["A1"] = "Top Attack Types"
    ws2["A1"].font = title_font
    ws2.append([])
    ws2.append(["Attack Type", "Count"])
    for cell in ws2[3]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    for atk in (report.attack_types or [])[:10]:
        ws2.append([atk.get("name", ""), atk.get("value", 0)])
        for cell in ws2[ws2.max_row]:
            cell.border = thin_border
            cell.fill = row_fill_1 if ws2.max_row % 2 == 0 else row_fill_2

    ws2.column_dimensions["A"].width = 30
    ws2.column_dimensions["B"].width = 15

    # Add chart
    chart = LineChart()
    chart.title = "Attack Distribution"
    chart.style = 10
    data = Reference(ws2, min_col=2, min_row=3, max_row=ws2.max_row)
    cats = Reference(ws2, min_col=1, min_row=4, max_row=ws2.max_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.shape = 4
    ws2.add_chart(chart, "E3")

    # Sheet 3: Attacker IPs
    ws3 = wb.create_sheet("Attacker IPs")
    ws3.sheet_properties.tabColor = "EF4444"
    ws3["A1"] = "Top Attacker IPs"
    ws3["A1"].font = title_font
    ws3.append([])
    ws3.append(["IP Address", "Country", "Attack Count", "Risk Level"])
    for cell in ws3[3]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    for ip_info in (report.attacker_ips or [])[:15]:
        ws3.append([
            ip_info.get("ip", ""),
            ip_info.get("country", "—"),
            ip_info.get("count", 0),
            ip_info.get("risk", ""),
        ])
        for cell in ws3[ws3.max_row]:
            cell.border = thin_border

    ws3.column_dimensions["A"].width = 22
    ws3.column_dimensions["B"].width = 15
    ws3.column_dimensions["C"].width = 15
    ws3.column_dimensions["D"].width = 15

    # Sheet 4: Timeline
    ws4 = wb.create_sheet("Timeline")
    ws4.sheet_properties.tabColor = "10B981"
    ws4["A1"] = "Attack Timeline"
    ws4["A1"].font = title_font
    ws4.append([])
    ws4.append(["Time", "Attacks"])
    for cell in ws4[3]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    for tl in (report.timeline_data or []):
        ws4.append([tl.get("time", ""), tl.get("attacks", 0)])
        for cell in ws4[ws4.max_row]:
            cell.border = thin_border

    ws4.column_dimensions["A"].width = 15
    ws4.column_dimensions["B"].width = 15

    # Timeline chart
    chart2 = LineChart()
    chart2.title = "Attack Timeline"
    chart2.x_axis.title = "Time"
    chart2.y_axis.title = "Attacks"
    data2 = Reference(ws4, min_col=2, min_row=3, max_row=ws4.max_row)
    cats2 = Reference(ws4, min_col=1, min_row=4, max_row=ws4.max_row)
    chart2.add_data(data2, titles_from_data=True)
    chart2.set_categories(cats2)
    ws4.add_chart(chart2, "D3")

    # Sheet 5: Subdomains (NEW)
    ws5 = wb.create_sheet("Subdomains")
    ws5.sheet_properties.tabColor = "8B5CF6"
    ws5["A1"] = "Top Subdomains Attacked"
    ws5["A1"].font = title_font
    ws5.append([])
    ws5.append(["Subdomain", "IP Address", "Attacks", "Country"])
    for cell in ws5[3]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="center")

    for sub in (report.subdomains or [])[:15]:
        ws5.append([
            sub.get("name", ""),
            sub.get("ip", "—"),
            sub.get("attacks", 0),
            sub.get("country", "—"),
        ])
        for cell in ws5[ws5.max_row]:
            cell.border = thin_border

    ws5.column_dimensions["A"].width = 35
    ws5.column_dimensions["B"].width = 18
    ws5.column_dimensions["C"].width = 12
    ws5.column_dimensions["D"].width = 15

    # Sheet 6: Conclusions & Recommendations (NEW)
    ws6 = wb.create_sheet("Conclusions")
    ws6.sheet_properties.tabColor = "EC4899"
    ws6["A1"] = "Kesimpulan & Rekomendasi"
    ws6["A1"].font = title_font

    # Conclusion
    ws6["A3"] = "Kesimpulan"
    ws6["A3"].font = Font(name="Calibri", size=13, bold=True, color="94A3B8")
    conclusion = narrative.get("conclusion", "")
    ws6.merge_cells("A4:D10")
    ws6["A4"] = conclusion
    ws6["A4"].font = Font(name="Calibri", size=10)
    ws6["A4"].alignment = Alignment(wrap_text=True, vertical="top")
    ws6.row_dimensions[4].height = 150

    # Recommendations
    ws6["A12"] = "Rekomendasi"
    ws6["A12"].font = Font(name="Calibri", size=13, bold=True, color="94A3B8")
    recommendations = narrative.get("recommendations", [])
    if isinstance(recommendations, list):
        for i, rec in enumerate(recommendations, start=13):
            ws6[f"A{i}"] = f"{i-12}. {rec}"
            ws6[f"A{i}"].font = Font(name="Calibri", size=10)
            ws6[f"A{i}"].alignment = Alignment(wrap_text=True)
            ws6.row_dimensions[i].height = 40

    ws6.column_dimensions["A"].width = 100

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    filename = f"FortiAnalis_Report_{datetime.now().strftime('%Y%m%d')}.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/export/{file_id}/word")
async def export_word(
    file_id: str,
    db: Session = Depends(get_db),
):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx is not installed.")

    report = db.query(Report).filter(Report.id == file_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")

    # Generate comprehensive narrative
    parsed_data = {
        "stats": {
            "totalRequests": report.total_requests,
            "totalAttacks": report.total_attacks,
            "blockedAttacks": report.blocked_attacks,
            "uniqueIps": report.unique_ips,
            "blockRate": report.block_rate,
        },
        "attackTypes": report.attack_types or [],
        "attackerIPs": report.attacker_ips or [],
        "subdomains": report.subdomains or [],
        "timelineData": report.timeline_data or [],
    }
    ai_insight = {
        "analysis": report.ai_analysis,
        "recommendation": report.ai_recommendation,
        "powered_by": report.ai_powered_by,
    }
    narrative = generate_report_narrative(parsed_data, ai_insight)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Title
    title = doc.add_heading("FortiAnalis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x06, 0xB6, 0xD4)

    doc.add_paragraph("WAF Security Report").runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    meta = doc.add_paragraph()
    meta.add_run(f"File: {report.filename}  |  Generated: {datetime.now().strftime('%d %b %Y %H:%M')}").font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("_" * 70)

    # Executive Summary (AI-generated)
    doc.add_heading("Executive Summary", level=1)
    exec_summary = narrative.get("executive_summary", "")
    for para in exec_summary.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Statistics Overview
    doc.add_heading("Statistik Keseluruhan", level=1)
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = "Light Shading Accent 1"
    summary_data = [
        ("Metric", "Value"),
        ("Total Requests", f"{report.total_requests:,}"),
        ("Total Attacks", f"{report.total_attacks:,}"),
        ("Blocked Attacks", f"{report.blocked_attacks:,}"),
        ("Block Rate", f"{report.block_rate}%"),
        ("Unique Attacker IPs", f"{report.unique_ips:,}"),
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    # Statistics narrative
    stats_narrative = narrative.get("statistics_narrative", "")
    for para in stats_narrative.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Attack Types
    if report.attack_types:
        doc.add_heading("Top Attack Types", level=1)
        atk_table = doc.add_table(rows=1, cols=2)
        atk_table.style = "Light Shading Accent 1"
        hdr = atk_table.rows[0]
        hdr.cells[0].text = "Attack Type"
        hdr.cells[1].text = "Count"
        for atk in (report.attack_types or [])[:10]:
            row = atk_table.add_row()
            row.cells[0].text = atk.get("name", "")
            row.cells[1].text = str(atk.get("value", 0))

        # Attack types narrative
        atk_narrative = narrative.get("attack_types_narrative", "")
        for para in atk_narrative.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Timeline narrative
    timeline_narrative = narrative.get("timeline_narrative", "")
    if timeline_narrative:
        doc.add_heading("Attack Timeline Analysis", level=1)
        for para in timeline_narrative.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Top Subdomains Attacked (NEW)
    if report.subdomains:
        doc.add_heading("Top Subdomains Attacked", level=1)
        subdomain_table = doc.add_table(rows=1, cols=4)
        subdomain_table.style = "Light Shading Accent 1"
        hdr = subdomain_table.rows[0]
        hdr.cells[0].text = "Subdomain"
        hdr.cells[1].text = "IP Address"
        hdr.cells[2].text = "Attacks"
        hdr.cells[3].text = "Country"
        for sub in (report.subdomains or [])[:10]:
            row = subdomain_table.add_row()
            row.cells[0].text = sub.get("name", "")
            row.cells[1].text = sub.get("ip", "—")
            row.cells[2].text = str(sub.get("attacks", 0))
            row.cells[3].text = sub.get("country", "—")

        # Subdomain narrative
        subdomain_narrative = narrative.get("subdomain_narrative", "")
        for para in subdomain_narrative.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Attacker IPs
    if report.attacker_ips:
        doc.add_heading("Top Attacker IPs", level=1)
        ip_table = doc.add_table(rows=1, cols=4)
        ip_table.style = "Light Shading Accent 1"
        hdr = ip_table.rows[0]
        hdr.cells[0].text = "IP Address"
        hdr.cells[1].text = "Country"
        hdr.cells[2].text = "Attack Count"
        hdr.cells[3].text = "Risk Level"
        for ip_info in (report.attacker_ips or [])[:15]:
            row = ip_table.add_row()
            row.cells[0].text = ip_info.get("ip", "")
            row.cells[1].text = ip_info.get("country", "—")
            row.cells[2].text = str(ip_info.get("count", 0))
            row.cells[3].text = ip_info.get("risk", "")

        # Attacker narrative
        attacker_narrative = narrative.get("attacker_narrative", "")
        for para in attacker_narrative.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # Conclusion (NEW)
    doc.add_heading("Kesimpulan", level=1)
    conclusion = narrative.get("conclusion", "")
    for para in conclusion.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # Recommendations (NEW)
    doc.add_heading("Rekomendasi", level=1)
    recommendations = narrative.get("recommendations", [])
    if isinstance(recommendations, list):
        for i, rec in enumerate(recommendations, 1):
            p = doc.add_paragraph(f"{i}. {rec}")
            p.style = "List Number"
    else:
        doc.add_paragraph(recommendations)

    # Footer
    doc.add_paragraph("_" * 70)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("FortiAnalis v2.0 — Confidential").font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"FortiAnalis_Report_{datetime.now().strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/export/{file_id}/csv")
async def export_csv(
    file_id: str,
    db: Session = Depends(get_db),
):
    report = db.query(Report).filter(Report.id == file_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")

    buffer = io.StringIO()
    writer = csv_module.writer(buffer)
    writer.writerow(["Metric", "Value"])
    writer.writerow(["Total Requests", report.total_requests])
    writer.writerow(["Total Attacks", report.total_attacks])
    writer.writerow(["Blocked Attacks", report.blocked_attacks])
    writer.writerow(["Block Rate", f"{report.block_rate}%"])
    writer.writerow(["Unique IPs", report.unique_ips])
    writer.writerow([])

    writer.writerow(["Attack Type", "Count"])
    for atk in (report.attack_types or []):
        writer.writerow([atk.get("name", ""), atk.get("value", 0)])
    writer.writerow([])

    writer.writerow(["Subdomain", "IP Address", "Attacks", "Country"])
    for sub in (report.subdomains or []):
        writer.writerow([sub.get("name", ""), sub.get("ip", "—"), sub.get("attacks", 0), sub.get("country", "—")])
    writer.writerow([])

    writer.writerow(["IP Address", "Country", "Attack Count", "Risk Level"])
    for ip_info in (report.attacker_ips or []):
        writer.writerow([ip_info.get("ip", ""), ip_info.get("country", "—"), ip_info.get("count", 0), ip_info.get("risk", "")])
    writer.writerow([])

    writer.writerow(["Time", "Attacks"])
    for tl in (report.timeline_data or []):
        writer.writerow([tl.get("time", ""), tl.get("attacks", 0)])

    csv_bytes = buffer.getvalue().encode("utf-8")
    filename = f"FortiAnalis_Report_{datetime.now().strftime('%Y%m%d')}.csv"
    return StreamingResponse(
        io.BytesIO(csv_bytes),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────────────────────────────────────────────────────────────────────────
#  PDF Generation with Charts (ReportLab Drawing API)
# ─────────────────────────────────────────────────────────────────────────────

def _generate_pdf(report: Report, period: str, narrative: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
            PageBreak,
        )
        from reportlab.graphics.shapes import Drawing, Rect, String
        from reportlab.graphics.charts.barcharts import VerticalBarChart
        from reportlab.graphics.charts.linecharts import HorizontalLineChart
        from reportlab.graphics.charts.piecharts import Pie
        from reportlab.lib.validators import Auto
    except ImportError:
        raise HTTPException(status_code=501, detail="reportlab is not installed.")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    primary_color = colors.HexColor("#06b6d4")
    dark_bg = colors.HexColor("#0f172a")

    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                  textColor=primary_color, fontSize=22, spaceAfter=6)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                               textColor=colors.HexColor("#94a3b8"), fontSize=13, spaceBefore=14)
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
                                 textColor=colors.HexColor("#334155"), fontSize=10, leading=15)
    small_style = ParagraphStyle("Small", parent=styles["Normal"],
                                  textColor=colors.HexColor("#64748b"), fontSize=9)

    now = datetime.now()
    if period == "daily":
        range_label = now.strftime("%d %B %Y")
    elif period == "weekly":
        week_start = now - timedelta(days=now.weekday())
        range_label = f"{week_start.strftime('%d %b')} – {now.strftime('%d %b %Y')}"
    else:
        range_label = now.strftime("%B %Y")

    elements = []

    # ── Header ────────────────────────────────────────────────────────────
    elements.append(Paragraph("FortiAnalis", title_style))
    elements.append(Paragraph(
        f"WAF Security Report — {period.capitalize()} | {range_label}",
        ParagraphStyle("Sub", parent=styles["Normal"],
                       textColor=colors.HexColor("#64748b"), fontSize=11, spaceAfter=4)
    ))
    elements.append(Paragraph(
        f"File: {report.filename} | Generated: {now.strftime('%d %b %Y %H:%M')}",
        small_style
    ))
    elements.append(HRFlowable(width="100%", thickness=1, color=primary_color, spaceAfter=12))

    # ── Executive Summary (AI-generated) ──────────────────────────────────
    elements.append(Paragraph("Executive Summary", h2_style))
    exec_summary = narrative.get("executive_summary", "")
    for para in exec_summary.split("\n\n"):
        if para.strip():
            elements.append(Paragraph(para.strip(), body_style))
            elements.append(Spacer(1, 8))
    elements.append(Spacer(1, 6))

    # ── Statistics Overview ───────────────────────────────────────────────
    elements.append(Paragraph("Statistik Keseluruhan", h2_style))
    stat_data = [
        ["Metric", "Value"],
        ["Total Requests", f"{report.total_requests:,}"],
        ["Total Attacks", f"{report.total_attacks:,}"],
        ["Blocked Attacks", f"{report.blocked_attacks:,}"],
        ["Block Rate", f"{report.block_rate}%"],
        ["Unique Attacker IPs", f"{report.unique_ips:,}"],
    ]
    stat_table = Table(stat_data, colWidths=[9*cm, 8*cm])
    stat_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), dark_bg),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, 0), 11),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
        ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#1e293b")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("FONTSIZE", (0, 1), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    elements.append(stat_table)
    elements.append(Spacer(1, 10))

    # Statistics narrative
    stats_narrative = narrative.get("statistics_narrative", "")
    for para in stats_narrative.split("\n\n"):
        if para.strip():
            elements.append(Paragraph(para.strip(), body_style))
            elements.append(Spacer(1, 6))
    elements.append(Spacer(1, 8))

    # ── Attack Types Bar Chart ────────────────────────────────────────────
    attack_types = report.attack_types or []
    if attack_types:
        elements.append(Paragraph("Top Attack Types", h2_style))

        drawing = Drawing(400, 200)
        bc = VerticalBarChart()
        bc.x = 50
        bc.y = 10
        bc.height = 150
        bc.width = 340
        bc.data = [[atk.get("value", 0) for atk in attack_types[:6]]]
        bc.strokeColor = colors.HexColor("#e2e8f0")
        bc.valueAxis.valueMin = 0
        bc.valueAxis.valueMax = max([a.get("value", 0) for a in attack_types[:6]] or [1]) * 1.2
        bc.valueAxis.valueStep = max(1, int(bc.valueAxis.valueMax / 5))
        bc.categoryAxis.labels.boxAnchor = "autox"
        bc.categoryAxis.categoryNames = [atk.get("name", "")[:12] for atk in attack_types[:6]]
        bc.categoryAxis.labels.fontSize = 8
        bc.bars[0].fillColor = primary_color
        drawing.add(bc)
        elements.append(drawing)
        elements.append(Spacer(1, 10))

        # Also include table for detail
        atk_data = [["Attack Type", "Count"]] + [
            [a["name"], str(a["value"])] for a in attack_types[:8]
        ]
        atk_table = Table(atk_data, colWidths=[12*cm, 5*cm])
        atk_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), dark_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
            ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#1e293b")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        elements.append(atk_table)
        elements.append(Spacer(1, 10))

        # Attack types narrative
        atk_narrative = narrative.get("attack_types_narrative", "")
        for para in atk_narrative.split("\n\n"):
            if para.strip():
                elements.append(Paragraph(para.strip(), body_style))
                elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 8))

    # ── Timeline Line Chart ───────────────────────────────────────────────
    timeline_data = report.timeline_data or []
    if timeline_data:
        elements.append(Paragraph("Attack Timeline", h2_style))

        drawing2 = Drawing(400, 180)
        lc = HorizontalLineChart()
        lc.x = 60
        lc.y = 10
        lc.height = 140
        lc.width = 330
        lc.data = [[tl.get("attacks", 0) for tl in timeline_data]]
        lc.joinedLines = True
        lc.lines[0].strokeColor = primary_color
        lc.lines[0].strokeWidth = 2
        lc.strokeColor = colors.HexColor("#e2e8f0")
        lc.valueAxis.valueMin = 0
        lc.valueAxis.valueMax = max([t.get("attacks", 0) for t in timeline_data] or [1]) * 1.2
        lc.valueAxis.valueStep = max(1, int(lc.valueAxis.valueMax / 5))
        lc.categoryAxis.categoryNames = [tl.get("time", "") for tl in timeline_data]
        lc.categoryAxis.labels.fontSize = 7
        lc.categoryAxis.labels.angle = 45
        drawing2.add(lc)
        elements.append(drawing2)
        elements.append(Spacer(1, 10))

        # Timeline narrative
        timeline_narrative = narrative.get("timeline_narrative", "")
        for para in timeline_narrative.split("\n\n"):
            if para.strip():
                elements.append(Paragraph(para.strip(), body_style))
                elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 8))

    # ── Top Subdomains Attacked (NEW) ─────────────────────────────────────
    subdomains = report.subdomains or []
    if subdomains:
        elements.append(Paragraph("Top Subdomains Attacked", h2_style))
        subdomain_data = [["Subdomain", "IP Address", "Attacks", "Country"]] + [
            [s.get("name", ""), s.get("ip", "—"), str(s.get("attacks", 0)), s.get("country", "—")]
            for s in subdomains[:10]
        ]
        subdomain_table = Table(subdomain_data, colWidths=[6*cm, 4*cm, 3*cm, 4*cm])
        subdomain_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), dark_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
            ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#1e293b")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        elements.append(subdomain_table)
        elements.append(Spacer(1, 10))

        # Subdomain narrative
        subdomain_narrative = narrative.get("subdomain_narrative", "")
        for para in subdomain_narrative.split("\n\n"):
            if para.strip():
                elements.append(Paragraph(para.strip(), body_style))
                elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 8))

    # ── Top Attacker IPs ──────────────────────────────────────────────────
    attacker_ips = report.attacker_ips or []
    if attacker_ips:
        elements.append(Paragraph("Top Attacker IPs", h2_style))
        ip_data = [["IP Address", "Country", "Attack Count", "Risk Level"]] + [
            [r.get("ip", ""), r.get("country", "—"), str(r.get("count", 0)), r.get("risk", "")]
            for r in attacker_ips[:10]
        ]
        ip_table = Table(ip_data, colWidths=[6*cm, 3*cm, 4*cm, 4*cm])
        ip_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), dark_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
            ("TEXTCOLOR", (0, 1), (-1, -1), colors.HexColor("#1e293b")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        elements.append(ip_table)
        elements.append(Spacer(1, 10))

        # Attacker narrative
        attacker_narrative = narrative.get("attacker_narrative", "")
        for para in attacker_narrative.split("\n\n"):
            if para.strip():
                elements.append(Paragraph(para.strip(), body_style))
                elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 8))

    # ── Conclusion (NEW) ──────────────────────────────────────────────────
    elements.append(Paragraph("Kesimpulan", h2_style))
    conclusion = narrative.get("conclusion", "")
    for para in conclusion.split("\n\n"):
        if para.strip():
            elements.append(Paragraph(para.strip(), body_style))
            elements.append(Spacer(1, 8))
    elements.append(Spacer(1, 10))

    # ── Recommendations (NEW) ─────────────────────────────────────────────
    elements.append(Paragraph("Rekomendasi", h2_style))
    recommendations = narrative.get("recommendations", [])
    if isinstance(recommendations, list):
        for i, rec in enumerate(recommendations, 1):
            elements.append(Paragraph(f"{i}. {rec}", body_style))
            elements.append(Spacer(1, 6))
    else:
        elements.append(Paragraph(recommendations, body_style))
    elements.append(Spacer(1, 10))

    # ── Footer ────────────────────────────────────────────────────────────
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#cbd5e1"), spaceAfter=6))
    elements.append(Paragraph(
        f"FortiAnalis v2.0 — Confidential | {now.strftime('%Y')}",
        small_style
    ))

    doc.build(elements)
    return buffer.getvalue()
